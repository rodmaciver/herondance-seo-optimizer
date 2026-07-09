"""Render the human-paste-ready output: markdown + Word, field-by-field."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo
from urllib.parse import urlparse

from docx import Document
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from .schema import ExecutionPlan, PageSnapshot

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"

DEV_MODE_WARNING = (
    "⚠ IMPORTANT: be very careful NOT to toggle Developer Mode while you are in "
    "this area. Only paste the mapping line into URL mappings and save."
)


def _slug_from_url(url: str) -> str:
    path = urlparse(url).path.strip("/")
    return path.replace("/", "-") or "home"


def _decision(plan: ExecutionPlan, field: str, default: str = "keep current") -> str:
    for item in plan.items:
        if item.field == field:
            return item.decision
    return default


def resolve_ads_final_url(plan: ExecutionPlan) -> tuple[str, bool]:
    """Return the planned final URL and whether a redirect is required."""
    url_slug = _decision(plan, "url_slug")
    parsed = urlparse(plan.page_url)
    base = f"{parsed.scheme}://{parsed.netloc}"
    if url_slug.lower().startswith("keep"):
        return plan.page_url.rstrip("/"), False

    final_url = f"{base}/{url_slug.lstrip('/')}"
    redirect = (plan.redirect_mapping or "").strip()
    is_changing = bool(redirect) and redirect.lower() not in ("null", "none", "n/a")
    return final_url, is_changing


def render_pastepack(
    plan: ExecutionPlan,
    snapshot: PageSnapshot,
    operator: str = "Reviewer",
    ad_assets: dict | None = None,
) -> dict:
    """Build the paste-pack and save both .md and .docx to output/.

    Returns {"markdown": str, "md_path": str, "docx_path": str}.
    ad_assets: optional dict with keys headlines (list[str]), descriptions (list[str]),
               flagged (bool), flag_reason (str). When None, the Google Ads section
               shows keyword terms only.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(ZoneInfo("America/New_York"))
    slug = _slug_from_url(plan.page_url)

    seo_title = _decision(plan, "seo_title")
    meta_description = _decision(plan, "meta_description")
    url_slug = _decision(plan, "url_slug")
    h1 = _decision(plan, "h1")

    ads_final_url, url_is_changing = resolve_ads_final_url(plan)

    # Dynamic section counter so numbers stay sequential when redirect is absent.
    sec = 1

    header_url = ads_final_url if url_is_changing else plan.page_url

    lines: list[str] = []
    lines.append(f"# SEO update: {header_url}")
    lines.append(f"Generated {timestamp:%Y-%m-%d %H:%M} · Reviewed and approved by {operator}")
    lines.append("")
    lines.append(f"## {sec}. Page settings → SEO tab")
    sec += 1
    # Trailing two spaces = Markdown hard line-break so Gradio renders each on its own line.
    lines.append(f"**SEO title:** {seo_title}  ")
    lines.append(f"**SEO description:** {meta_description}")
    lines.append("")
    lines.append(f"## {sec}. Page settings → General tab")
    sec += 1
    lines.append(f"**URL slug:** {url_slug}  ")
    lines.append(f"**Page title:** {h1}")
    lines.append("")

    # Redirect section removed — mapping line is written to column C of the queue sheet
    # and the new URL is shown in the header above instead.

    lines.append(f"## {sec}. Page content changes (in the Squarespace editor)")
    sec += 1
    for i, bc in enumerate(plan.body_changes, start=1):
        action_label = bc.action.replace("remove_section", "remove").replace("add_section", "add")
        lines.append(f"{i}. [{action_label}]  ")
        lines.append(f"   **Instruction:** {bc.instruction}  ")
        if bc.new_text:
            lines.append(f"   **New text:** {bc.new_text}")
    lines.append("")

    # --- Google Ads section removed from paste-pack (now in separate spreadsheet) ---
    # lines.append(f"## {sec}. Google Ads")
    # sec += 1
    # lines.append(f"**Final URL:** {ads_final_url}  ")
    # if url_is_changing:
    #     lines.append(
    #         "⚠ Create campaign in PAUSE mode — can be prepared now, "
    #         "activate only after URL redirect is live in Squarespace.  "
    #     )
    # lines.append("")
    # if ad_assets:
    #     if ad_assets.get("flagged"):
    #         lines.append(f"⚠ MANUAL REVIEW REQUIRED: {ad_assets.get('flag_reason', '')}")
    #         lines.append("")
    #     if ad_assets.get("campaign_slug"):
    #         lines.append(f"**Campaign:** {ad_assets['campaign_slug']}")
    #     if ad_assets.get("ad_group_slug"):
    #         lines.append(f"**Ad group:** {ad_assets['ad_group_slug']}")
    #     lines.append("")
    #     lines.append("**Responsive Search Ad —**  ")
    #     lines.append("**Headlines:**  ")
    #     for i, h in enumerate(ad_assets.get("headlines", []), 1):
    #         lines.append(f"{i}. {h}  ")
    #     lines.append("")
    #     lines.append("**Descriptions:**  ")
    #     for i, d in enumerate(ad_assets.get("descriptions", []), 1):
    #         lines.append(f"{i}. {d}  ")
    #     lines.append("")
    #     core_kws = ad_assets.get("core_keywords", [])
    #     variants_map = ad_assets.get("keyword_variants", {})
    #     if core_kws:
    #         lines.append("**Core keywords and variants** — Phrase Match, Manual CPC, Max CPC $0.20:  ")
    #         for kw in core_kws:
    #             kw_clean = kw.strip("[]")
    #             lines.append(f'- "{kw_clean}"  ')
    #             key = kw if kw in variants_map else kw_clean
    #             for v in variants_map.get(key, []):
    #                 lines.append(f'  - "{v.strip("[]")}"  ')
    #         lines.append("")
    #     neg_kws = ad_assets.get("negative_keywords", [])
    #     if neg_kws:
    #         lines.append("**Page-specific negative keywords:**  ")
    #         for nk in neg_kws:
    #             lines.append(f"- {nk}  ")
    #         lines.append("")
    # else:
    #     lines.append("*Google Ads copy unavailable for this page.*  ")
    #     lines.append("")

    markdown = "\n".join(lines)

    md_path = OUTPUT_DIR / f"{slug}_{timestamp:%Y%m%d}.md"
    md_path.write_text(markdown)

    docx_path = OUTPUT_DIR / f"{slug}_{timestamp:%Y%m%d}.docx"
    _write_docx(docx_path, plan, snapshot, operator, timestamp,
                seo_title, meta_description, url_slug, h1, ad_assets,
                ads_final_url, url_is_changing)

    return {"markdown": markdown, "md_path": str(md_path), "docx_path": str(docx_path),
            "ads_final_url": ads_final_url}


_LABEL_COLOR = RGBColor(0x2E, 0x74, 0xB5)  # Word standard blue


def _lp(doc: Document, label: str, value: str = "", space_before: bool = False) -> Paragraph:
    """Add a paragraph with a colored bold label and plain value text."""
    p = doc.add_paragraph()
    if space_before:
        p.paragraph_format.space_before = Pt(6)
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = _LABEL_COLOR
    if value:
        p.add_run(value)
    return p


def _write_docx(
    path: Path,
    plan: ExecutionPlan,
    snapshot: PageSnapshot,
    operator: str,
    timestamp: datetime,
    seo_title: str,
    meta_description: str,
    url_slug: str,
    h1: str,
    ad_assets: dict | None = None,
    ads_final_url: str = "",
    url_is_changing: bool = False,
) -> None:
    sec = 1

    header_url = ads_final_url if url_is_changing else plan.page_url

    doc = Document()
    doc.add_heading(f"SEO update: {header_url}", level=1)
    doc.add_paragraph(f"Generated {timestamp:%Y-%m-%d %H:%M} · Reviewed and approved by {operator}")

    doc.add_heading(f"{sec}. Page settings → SEO tab", level=2)
    sec += 1
    _lp(doc, "SEO title: ", seo_title)
    _lp(doc, "SEO description: ", meta_description)

    doc.add_heading(f"{sec}. Page settings → General tab", level=2)
    sec += 1
    _lp(doc, "URL slug: ", url_slug)
    _lp(doc, "Page title: ", h1)

    # Redirect section removed — mapping line is written to column C of the queue sheet.

    doc.add_heading(f"{sec}. Page content changes (in the Squarespace editor)", level=2)
    sec += 1
    for i, bc in enumerate(plan.body_changes, start=1):
        action_label = bc.action.replace("remove_section", "remove").replace("add_section", "add")
        doc.add_paragraph(f"[{action_label}]", style="List Number")
        _lp(doc, "Instruction: ", bc.instruction)
        if bc.new_text:
            _lp(doc, "New text: ", bc.new_text)

    # --- Google Ads section removed from docx (now in separate spreadsheet) ---
    # doc.add_heading(f"{sec}. Google Ads", level=2)
    # sec += 1
    # _lp(doc, "Final URL: ", ads_final_url)
    # if url_is_changing:
    #     p = doc.add_paragraph()
    #     run = p.add_run(
    #         "⚠ Create campaign in PAUSE mode — can be prepared now, "
    #         "activate only after URL redirect is live in Squarespace."
    #     )
    #     run.bold = True
    # if ad_assets:
    #     if ad_assets.get("flagged"):
    #         p = doc.add_paragraph()
    #         run = p.add_run(f"⚠ MANUAL REVIEW REQUIRED: {ad_assets.get('flag_reason', '')}")
    #         run.bold = True
    #     if ad_assets.get("campaign_slug"):
    #         _lp(doc, "Campaign: ", ad_assets["campaign_slug"])
    #     if ad_assets.get("ad_group_slug"):
    #         _lp(doc, "Ad group: ", ad_assets["ad_group_slug"])
    #     rsa_label = _lp(doc, "Responsive Search Ad —", space_before=True)
    #     rsa_label.paragraph_format.keep_with_next = True
    #     headlines_label = _lp(doc, "Headlines:")
    #     headlines_label.paragraph_format.keep_with_next = True
    #     headlines = ad_assets.get("headlines", [])
    #     for i, h in enumerate(headlines, 1):
    #         p = doc.add_paragraph(f"{i}. {h}")
    #         if i < len(headlines):
    #             p.paragraph_format.keep_with_next = True
    #     descriptions_label = _lp(doc, "Descriptions:", space_before=True)
    #     descriptions_label.paragraph_format.keep_with_next = True
    #     descriptions = ad_assets.get("descriptions", [])
    #     for i, d in enumerate(descriptions, 1):
    #         p = doc.add_paragraph(f"{i}. {d}")
    #         if i < len(descriptions):
    #             p.paragraph_format.keep_with_next = True
    #     core_kws = ad_assets.get("core_keywords", [])
    #     variants_map = ad_assets.get("keyword_variants", {})
    #     if core_kws:
    #         _lp(doc, "Core keywords and variants — Phrase Match, Manual CPC, Max CPC $0.20:", space_before=True)
    #         for kw in core_kws:
    #             kw_clean = kw.strip("[]")
    #             doc.add_paragraph(f'"{kw_clean}"', style="List Bullet")
    #             key = kw if kw in variants_map else kw_clean
    #             for v in variants_map.get(key, []):
    #                 p = doc.add_paragraph(style="List Bullet 2")
    #                 p.add_run(f'"{v.strip("[]")}"')
    #     neg_kws = ad_assets.get("negative_keywords", [])
    #     if neg_kws:
    #         _lp(doc, "Page-specific negative keywords:", space_before=True)
    #         for nk in neg_kws:
    #             doc.add_paragraph(nk, style="List Bullet")
    # else:
    #     doc.add_paragraph("Google Ads copy unavailable for this page.")

    doc.save(path)
