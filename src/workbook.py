"""Work-queue and registry I/O against the client's SEO workbook.

The workbook in ``data/`` is the client's own file and is treated as
read-only: every mutation operates on a timestamped copy in ``output/``.
"""
from __future__ import annotations

import logging
import re
import shutil
from datetime import datetime
from pathlib import Path
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse

import numpy as np
import openpyxl
import pandas as pd
from docx import Document

QUEUE_SHEET = "URLs to Do"
DONE_SHEET = "URLs Done"
QUEUE_COLUMNS = ["url", "clicks", "impressions", "ctr", "avg_position"]

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"

BASE_HOST = "https://herondance.org"

# Params that indicate tracking/analytics variants of the same page.
_TRACKING_PARAMS = frozenset({
    "utm_source", "utm_medium", "utm_campaign", "utm_term", "utm_content",
    "fbclid", "gclid", "_ga",
})

# CTR-by-position curve (client-provided reference points), interpolated.
_CTR_POSITIONS = [1, 2, 3, 5, 10, 20, 50, 100]
_CTR_VALUES = [0.28, 0.15, 0.10, 0.06, 0.025, 0.01, 0.004, 0.001]


# ---------------------------------------------------------------------------
# URL normalization
# ---------------------------------------------------------------------------

def normalize_url(raw: str, base_host: str = BASE_HOST) -> str:
    """Return the canonical form of a URL.

    Handles: bare paths, double slashes in the path, trailing slashes,
    HTML-encoded ampersands, and tracking-only query parameters.
    Applied consistently everywhere a URL is read, compared, fetched, or
    written back.
    """
    raw = str(raw).strip()
    if not raw or raw.lower() in ("nan", "none", ""):
        return ""

    # Bare path → prepend host.
    if raw.startswith("/"):
        raw = base_host.rstrip("/") + raw
    elif "://" not in raw:
        raw = base_host.rstrip("/") + "/" + raw.lstrip("/")

    parsed = urlparse(raw)

    # Collapse accidental double slashes in path.
    path = re.sub(r"/{2,}", "/", parsed.path)

    # Strip trailing slash (keep bare root "/").
    path = path.rstrip("/") or "/"

    # Decode HTML-encoded ampersands, then strip tracking params.
    query_str = parsed.query.replace("&amp;", "&")
    params = [
        (k, v) for k, v in parse_qsl(query_str, keep_blank_values=True)
        if k.lower() not in _TRACKING_PARAMS
    ]

    return urlunparse((parsed.scheme, parsed.netloc, path, "", urlencode(params), ""))


# ---------------------------------------------------------------------------
# Opportunity scoring
# ---------------------------------------------------------------------------

def expected_ctr(position: float) -> float:
    """Expected CTR for a given average search position (interpolated curve)."""
    position = max(1.0, float(position))
    return float(np.interp(position, _CTR_POSITIONS, _CTR_VALUES))


def position_factor(position: float) -> float:
    """0..1 weight peaking for pages near-but-not-on page one (positions 4-20)."""
    position = float(position)
    if position <= 1:
        return 0.2
    if position < 4:
        return 0.2 + (position - 1) / (4 - 1) * 0.8
    if position <= 20:
        return 1.0
    if position < 50:
        return 1.0 - (position - 20) / (50 - 20)
    return 0.0


def opportunity_score(df: pd.DataFrame) -> pd.DataFrame:
    """Add `score` and `why` columns, sorted by score descending (NaN last)."""
    df = df.copy()

    def _score(row: pd.Series) -> float:
        if pd.isna(row["impressions"]) or pd.isna(row["avg_position"]) or pd.isna(row["ctr"]):
            return float("nan")
        pf = position_factor(row["avg_position"])
        gap = max(0.0, expected_ctr(row["avg_position"]) - row["ctr"])
        return row["impressions"] * pf * gap

    def _why(row: pd.Series) -> str:
        if pd.isna(row.get("score")):
            return "No search data."
        return (
            f"{row['impressions']:,.0f} impressions at position "
            f"{row['avg_position']:.1f} — large page-one opportunity."
        )

    df["score"] = df.apply(_score, axis=1)
    df["why"] = df.apply(_why, axis=1)
    df = df.sort_values("score", ascending=False, na_position="last").reset_index(drop=True)
    return df


# ---------------------------------------------------------------------------
# Queue loading — two-section parsing
# ---------------------------------------------------------------------------

def _parse_section(raw: pd.DataFrame) -> pd.DataFrame:
    """Parse a raw slice of the workbook sheet into a clean queue DataFrame.

    Filters to rows with valid (http) URLs after normalization; metric
    columns are kept as NaN where blank (not coerced to 0).

    Extra columns added for batch tracking:
      _excel_row  — 1-based row number in the source xlsx (for write-back)
      generated   — True if column B contains "Generated"
    """
    if raw.empty:
        return pd.DataFrame(columns=QUEUE_COLUMNS + ["generated", "_excel_row"])

    section = raw.iloc[:, :5].copy()
    while len(section.columns) < 5:
        section[len(section.columns)] = None

    # Read column B (index 1) for "Generated" marker BEFORE renaming.
    # .astype(bool) ensures numpy bool dtype so ~ works correctly.
    # pd.Series(..., index=...) ensures index-aligned assignment survives the URL filter.
    generated = section.iloc[:, 1].apply(
        lambda x: str(x).strip().lower() == "generated" if pd.notna(x) else False
    ).astype(bool)
    excel_rows = pd.Series(section.index + 1, index=section.index)  # 1-based, index-aligned

    section.columns = QUEUE_COLUMNS

    section["url"] = section["url"].apply(
        lambda x: normalize_url(str(x)) if pd.notna(x) else ""
    )
    # Drop header rows, blank rows, and non-URL rows.
    section = section[section["url"].str.startswith("http")].copy()

    for col in ["clicks", "impressions", "ctr", "avg_position"]:
        section[col] = pd.to_numeric(section[col], errors="coerce")

    section["generated"] = generated
    section["_excel_row"] = excel_rows

    return section.reset_index(drop=True)


def parse_queue_raw(raw: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Split a raw DataFrame (no headers) into (priority_df, backlog_df).

    Shared by both the xlsx path and the Google Sheets path so parsing
    logic stays in one place.
    """
    header_positions = [
        i for i, val in enumerate(raw.iloc[:, 0])
        if str(val).strip().lower() == "top pages"
    ]

    divider_pos: int | None = None
    if len(header_positions) >= 2:
        divider_pos = header_positions[1]
    elif len(header_positions) == 1:
        if header_positions[0] > 0:
            divider_pos = header_positions[0]
        else:
            logging.getLogger(__name__).warning(
                "Queue sheet: 'Top Pages' header found only at row 0 — "
                "expected a second occurrence as a backlog divider. "
                "Treating all rows as backlog; check the sheet structure."
            )

    if divider_pos is not None:
        priority_raw = raw.iloc[:divider_pos]
        backlog_raw = raw.iloc[divider_pos + 1:]
    else:
        priority_raw = pd.DataFrame()
        backlog_raw = raw

    return _parse_section(priority_raw), _parse_section(backlog_raw)


def load_queue(path: str | Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Read 'URLs to Do', detect the priority/backlog divider, return
    (priority_df, backlog_df).

    The divider is the repeated header row whose first cell reads "Top pages"
    (case-insensitive).  Everything above it is the priority list; everything
    below is the full backlog.

    - If no divider exists the whole sheet is returned as the backlog with an
      empty priority DataFrame.
    - Metric columns may be NaN for pages with no search data.
    - All URLs are normalized before return.
    """
    raw = pd.read_excel(path, sheet_name=QUEUE_SHEET, header=None)
    return parse_queue_raw(raw)


# ---------------------------------------------------------------------------
# Workbook mutations
# ---------------------------------------------------------------------------

def mark_done(
    path: str | Path,
    old_url: str,
    new_url: str,
    new_title: str,
    stats: dict,
) -> str:
    """Copy the workbook, move the page from 'URLs to Do' to 'URLs Done'.

    `stats` is a dict with keys clicks, impressions, ctr, avg_position.
    Returns the path to the new copy. The original file is never mutated.
    URL matching uses normalization so variants of the same page still match.
    """
    path = Path(path)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    new_path = OUTPUT_DIR / f"seo_workbook_updated_{timestamp}.xlsx"
    shutil.copy(path, new_path)

    wb = openpyxl.load_workbook(new_path)

    # Remove the row for old_url from "URLs to Do" (normalize both sides).
    todo_ws = wb[QUEUE_SHEET]
    normalized_old = normalize_url(old_url)
    found = False
    for row in range(todo_ws.max_row, 0, -1):
        cell_val = todo_ws.cell(row=row, column=1).value
        if cell_val and normalize_url(str(cell_val)) == normalized_old:
            todo_ws.delete_rows(row, 1)
            found = True
            break
    if not found:
        logging.getLogger(__name__).warning(
            "mark_done: '%s' not found in '%s' sheet — appending to done without removing from queue.",
            old_url,
            QUEUE_SHEET,
        )

    # Append to "URLs Done": Former URL | 301 Redirect to | New page title | stats…
    done_ws = wb[DONE_SHEET]
    next_row = done_ws.max_row + 1
    while done_ws.cell(row=next_row - 1, column=1).value is None and next_row > 2:
        next_row -= 1
    done_ws.cell(row=next_row, column=1, value=old_url)
    done_ws.cell(row=next_row, column=2, value=new_url)
    done_ws.cell(row=next_row, column=3, value=new_title)
    done_ws.cell(row=next_row, column=4, value=stats.get("clicks"))
    done_ws.cell(row=next_row, column=5, value=stats.get("impressions"))
    done_ws.cell(row=next_row, column=6, value=stats.get("ctr"))
    done_ws.cell(row=next_row, column=7, value=stats.get("avg_position"))

    wb.save(new_path)
    return str(new_path)


def export_master_list(path: str | Path) -> str:
    """Write output/master_urls_and_titles.xlsx and .docx from 'URLs Done'."""
    path = Path(path)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    wb = openpyxl.load_workbook(path)
    done_ws = wb[DONE_SHEET]

    rows = []
    for row in done_ws.iter_rows(min_row=2, values_only=True):
        url, redirect_to = row[0], row[1]
        title = row[2] if len(row) > 2 else None
        if not url:
            continue
        current_url = redirect_to or url
        rows.append((current_url, str(title) if title is not None else ""))

    # .xlsx
    out_wb = openpyxl.Workbook()
    out_ws = out_wb.active
    out_ws.title = "Master List"
    out_ws.append(["URL", "Page Title"])
    for current_url, title in rows:
        out_ws.append([current_url, title])
    xlsx_path = OUTPUT_DIR / "master_urls_and_titles.xlsx"
    out_wb.save(xlsx_path)

    # .docx
    doc = Document()
    doc.add_heading("Master URL and Page Title List", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "URL"
    hdr[1].text = "Page Title"
    for current_url, title in rows:
        cells = table.add_row().cells
        cells[0].text = current_url
        cells[1].text = title
    docx_path = OUTPUT_DIR / "master_urls_and_titles.docx"
    doc.save(docx_path)

    return str(xlsx_path)
